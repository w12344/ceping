from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/chenpan/Documents/于无声处起惊雷/output/小凡教育科技业务流程与运营制度手册-第一版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"


def set_east_asia(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_para(p, before=0, after=6, line=1.25, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_run(p, text, size=11, bold=False, color=INK):
    r = p.add_run(text)
    set_east_asia(r)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return r


def add_p(doc, text="", size=11, bold=False, color=INK, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align)
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None, color=INK, size=10.5):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.15)
    for idx, part in enumerate(str(text).split("\n")):
        if idx:
            p.add_run().add_break()
        add_run(p, part, size=size, bold=bold, color=color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.tcW
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill=header_fill, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para(p, before=18, after=10, line=1.15)
        add_run(p, text, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_para(p, before=14, after=7, line=1.15)
        add_run(p, text, size=13, bold=True, color=BLUE)
    else:
        set_para(p, before=10, after=5, line=1.15)
        add_run(p, text, size=12, bold=True, color=DARK_BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=4, line=1.25)
        add_run(p, item, size=11)


def add_steps(doc, steps):
    for idx, (title, body) in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        set_para(p, after=4, line=1.25)
        add_run(p, f"{title}：", bold=True)
        add_run(p, body)


def add_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_after = Pt(12)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    set_para(p, after=3, line=1.2)
    add_run(p, title, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line=1.2)
    add_run(p2, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = FONT
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        styles[style_name].font.size = Pt(11)
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    set_para(header, after=0, line=1.0)
    add_run(header, "小凡教育科技业务流程与运营制度手册", size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    set_para(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(footer, "内部运营参考文件 · 第一版", size=9, color=MUTED)
    return doc


def build():
    doc = setup_document()

    add_p(doc, "小凡教育科技", size=12, bold=True, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "业务流程与运营制度手册", size=26, bold=True, color=INK, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "第一版底稿｜用于业务协同、岗位交接、运营管理与新人培训", size=13, color=MUTED, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rule(doc)
    add_callout(doc, "手册定位", "本手册用于把公司的业务链路、岗位协同、服务标准、管理制度和复盘机制沉淀为统一语言。它不是静态制度汇编，而是一份可以被培训、执行、检查和迭代的运营参考文件。")
    add_table(
        doc,
        ["版本", "适用范围", "主要使用者", "更新建议"],
        [["第一版底稿", "集团及各业务品牌通用流程", "管理者、校区负责人、咨询、教学、运营、教务、市场、HR、AI与系统支持岗位", "建议每季度结合业务复盘更新一次"]],
        [1350, 2350, 3900, 1760],
        header_fill=LIGHT_BLUE,
    )
    doc.add_page_break()

    add_heading(doc, "目录式导读", 1)
    add_p(doc, "建议将本手册作为组织运行的“操作系统说明书”。阅读时不必从头到尾逐字阅读，管理者可先看全景和制度原则，一线岗位可直接进入与自己相关的流程章节。")
    add_table(
        doc,
        ["模块", "回答的问题", "适合谁重点阅读"],
        [
            ["一、手册导读", "这本手册怎么用、按什么原则执行", "全员"],
            ["二、业务全景", "公司业务如何分工、如何协同", "管理层、跨部门岗位、新人"],
            ["三、用户全生命周期", "一个用户从线索到长期关系如何被服务", "市场、咨询、教务、教学、运营"],
            ["四、教学与服务流程", "课程、学情、反馈、家校沟通如何标准化", "教学、教研、教务、班主任、校区负责人"],
            ["五、销售与市场流程", "线索、跟进、转化、复盘如何闭环", "市场、咨询、运营、校区负责人"],
            ["六、运营管理制度", "合同、收款、排课、退款、客诉等如何执行", "运营、教务、财务协同、校区负责人"],
            ["七、人员协作机制", "岗位如何分工，跨部门如何交接与升级", "管理者、HR、项目负责人"],
            ["八、质量管理与复盘", "如何检查质量、发现问题、推动改进", "管理层、业务负责人、运营负责人"],
            ["九、AI与系统规范", "1605与数据系统如何支撑业务", "全员，重点为教学、运营、产品与AI支持"],
            ["十、附录工具包", "哪些模板可直接拿来用", "所有执行岗位"],
        ],
        [1300, 4750, 3310],
    )

    add_heading(doc, "一、手册导读", 1)
    add_heading(doc, "1.1 编写原则", 2)
    add_bullets(doc, [
        "以用户体验为中心：任何流程设计都要服务学生成长、家长信任和业务长期价值。",
        "以岗位责任为边界：每个关键动作必须有明确责任人，避免“大家都知道，但没人负责”。",
        "以数据和事实为依据：咨询、教学、服务、复盘都要尽量留下可追踪记录。",
        "以闭环为标准：流程不是完成动作，而是让问题被发现、被响应、被复盘、被改进。",
        "以简洁可执行为优先：制度应帮助一线更快做对事，而不是制造额外负担。",
    ])
    add_heading(doc, "1.2 使用方式", 2)
    add_steps(doc, [
        ("新人学习", "先阅读业务全景、用户生命周期和本岗位相关流程，理解公司如何从线索到服务形成闭环。"),
        ("岗位交接", "按流程清单核对客户、课程、数据、合同、家校沟通、异常事项和待办任务。"),
        ("管理检查", "围绕关键节点检查记录是否完整、责任是否明确、风险是否升级、结果是否复盘。"),
        ("制度迭代", "每季度收集一线反馈，将高频问题转化为流程补充、模板优化或系统字段更新。"),
    ])
    add_callout(doc, "执行口径", "当实际情况与手册不一致时，先保证用户利益和风险可控，再由责任岗位记录偏差原因并发起流程更新。")

    add_heading(doc, "二、公司业务全景", 1)
    add_heading(doc, "2.1 业务结构", 2)
    add_p(doc, "小凡教育科技的业务可理解为“教育品牌 + 技术底座 + 组织能力”的组合。各业务品牌面向不同教育场景，1605人工智能作为技术底座，为教学、运营、学情分析、内容生产和管理决策提供支持。")
    add_table(
        doc,
        ["业务单元", "定位", "核心任务", "关键协同"],
        [
            ["非凡教育", "面向艺考等升学场景的教育品牌", "围绕学生阶段性目标，提供课程、陪伴、反馈与升学支持", "市场获客、咨询转化、教学服务、活动运营、AI学情支持"],
            ["小凡私塾", "小而精的学生成长学院", "以小班、个性化、大师课等配置陪伴学生成长", "家校沟通、学情跟踪、学生状态管理、成长活动"],
            ["小凡公学", "面向未来教育探索的业务板块", "沉淀课程体系、学习方式与教育实验机制", "教研、项目制学习、专家资源、内容产品"],
            ["1605人工智能", "技术底座与智能支持系统", "提升数据、内容、学情、运营和管理效率", "与各业务系统、教学场景、运营流程深度连接"],
        ],
        [1500, 2450, 3000, 2410],
    )
    add_heading(doc, "2.2 业务运行总链路", 2)
    add_steps(doc, [
        ("被看见", "通过品牌内容、活动、渠道、转介绍和社群运营，让目标用户知道小凡教育科技。"),
        ("被理解", "通过咨询、诊断、测评和沟通，理解学生基础、目标、状态与家庭期待。"),
        ("被匹配", "基于诊断结果匹配课程、班型、老师、服务节奏和成长方案。"),
        ("被陪伴", "通过教学、作业、反馈、家校沟通和阶段复盘持续陪伴学生。"),
        ("被验证", "用学习结果、状态变化、家长反馈、续费转介绍等指标验证服务价值。"),
        ("被沉淀", "将优秀案例、教学经验、流程方法和数据洞察沉淀为组织资产。"),
    ])

    add_heading(doc, "三、用户全生命周期流程", 1)
    add_p(doc, "用户全生命周期流程是本手册的主流程。所有部门的工作最终都应回到同一条链路：让用户被准确理解、被合适服务、被持续反馈、被长期维护。")
    add_table(
        doc,
        ["阶段", "关键目标", "责任主岗", "必须留下的记录"],
        [
            ["线索获取", "识别有效用户来源", "市场/运营", "来源渠道、用户标签、首次触点、需求关键词"],
            ["咨询接待", "建立信任并完成基础信息收集", "咨询", "学生信息、家长诉求、目标、时间线、预算与顾虑"],
            ["诊断评估", "判断基础、问题和机会点", "咨询/教学/测评支持", "诊断表、测评结果、学习状态描述"],
            ["方案制定", "形成可执行学习方案", "咨询/教学负责人", "推荐课程、班型、师资、服务节奏、预期目标"],
            ["报名转化", "完成合同、缴费和入学确认", "咨询/运营/财务协同", "合同、缴费凭证、班级安排、交接单"],
            ["入学交接", "让教学服务准确接住用户", "咨询/教务/班主任", "学生画像、承诺事项、风险点、家长偏好"],
            ["学习服务", "持续交付课程与陪伴", "教学/教务/班主任", "出勤、作业、反馈、阶段测评、沟通记录"],
            ["阶段复盘", "判断效果并调整策略", "教学负责人/班主任", "阶段报告、问题清单、调整方案"],
            ["续费转介绍", "延续信任并放大口碑", "咨询/班主任/运营", "满意度、续费意向、转介绍线索"],
            ["结课维护", "保持长期关系和品牌资产", "运营/班主任", "结课报告、祝福触点、后续服务机会"],
        ],
        [1250, 2500, 1900, 3710],
    )
    add_heading(doc, "3.1 咨询接待流程", 2)
    add_steps(doc, [
        ("确认来源", "记录用户从哪里来，区分广告、活动、转介绍、自然咨询、社群等来源。"),
        ("建立信任", "先回应用户最关心的问题，不急于推产品，避免未经诊断直接承诺结果。"),
        ("收集信息", "围绕学生年级、基础、目标、时间、状态、过往学习经历和家庭期待建立初步画像。"),
        ("识别需求", "区分显性需求和真实问题，例如升学焦虑背后可能是基础薄弱、方法失效或状态不稳定。"),
        ("安排下一步", "根据情况安排测评、试听、诊断沟通或学习方案说明。"),
    ])
    add_heading(doc, "3.2 诊断评估流程", 2)
    add_bullets(doc, [
        "诊断不是为了证明学生“不行”，而是为了看见学生真实起点和可改变的路径。",
        "诊断结果应包括基础水平、学习习惯、目标差距、状态风险和可优先突破的关键问题。",
        "涉及专业课、文化课、心理状态或家庭沟通的复杂问题，应邀请对应负责人参与判断。",
        "诊断结论必须转化为方案，不停留在笼统评价。"
    ])
    add_heading(doc, "3.3 入学交接流程", 2)
    add_steps(doc, [
        ("咨询提交交接单", "包含学生基本情况、家长核心诉求、已承诺事项、敏感点、风险点和推荐方案。"),
        ("教务确认排课", "确认班型、老师、时间、教室/线上入口、教材资料和首次上课提醒。"),
        ("教学负责人复核", "确认课程匹配度和教师接手重点，必要时调整方案。"),
        ("班主任建立关系", "完成入班欢迎、家长沟通群建立、规则说明和首次反馈时间约定。"),
        ("首周回访", "入学后一周内确认适应情况，发现问题及时调整。"),
    ])

    add_heading(doc, "四、教学与服务流程", 1)
    add_heading(doc, "4.1 教学服务基本闭环", 2)
    add_steps(doc, [
        ("课前准备", "教师明确本节课目标、学生基础差异、重点难点、课堂产出和作业要求。"),
        ("课堂交付", "课堂应有清晰目标、有效互动、即时反馈和可被学生带走的方法或成果。"),
        ("课后记录", "教师完成出勤、课堂表现、掌握情况、作业布置和需要跟进的问题记录。"),
        ("学情反馈", "班主任或教师按约定节奏向家长反馈真实进展，既讲结果，也讲原因和下一步。"),
        ("阶段复盘", "按阶段输出学习报告，判断目标达成情况，并调整课程、作业或陪伴策略。"),
    ])
    add_heading(doc, "4.2 教研备课标准", 2)
    add_bullets(doc, [
        "课程目标要具体：学生学完后应能说清、做出或改正什么。",
        "课件和材料要可复用：优秀内容进入教研资产库，而不是只停留在个人电脑里。",
        "备课要考虑分层：同一班级内基础不同的学生，应有基础任务、提升任务和挑战任务。",
        "每次重要课程后，教师应记录一个可复盘问题和一个可沉淀经验。",
    ])
    add_heading(doc, "4.3 家校沟通标准", 2)
    add_table(
        doc,
        ["沟通场景", "沟通重点", "建议频率", "注意事项"],
        [
            ["入学初期", "规则、目标、服务方式、反馈节奏", "入学后1周内", "避免承诺无法完全控制的结果"],
            ["日常学习", "出勤、作业、课堂状态、近期改进", "按班型约定", "真实具体，少用空泛评价"],
            ["阶段复盘", "目标差距、进步证据、问题原因、调整动作", "每阶段一次", "要有下一步方案"],
            ["异常情况", "缺勤、情绪波动、成绩下滑、家长投诉", "发现后及时", "先安抚，再查因，最后给方案"],
            ["结课节点", "学习成果、成长变化、后续建议", "结课前后", "维护长期关系，沉淀口碑"],
        ],
        [1400, 3100, 1500, 3360],
    )

    add_heading(doc, "五、销售与市场流程", 1)
    add_heading(doc, "5.1 市场活动流程", 2)
    add_steps(doc, [
        ("明确目标", "区分品牌曝光、线索收集、转化促进、老用户维护或转介绍激活。"),
        ("设计内容", "活动主题必须与用户真实需求相关，避免为了热闹而热闹。"),
        ("准备物料", "包括海报、话术、报名表、签到表、回访标签、转化路径和风险预案。"),
        ("现场执行", "明确主持、接待、咨询、拍摄、数据记录和应急负责人。"),
        ("活动复盘", "活动结束后复盘到访、留资、有效线索、转化、成本、用户反馈和可复制经验。"),
    ])
    add_heading(doc, "5.2 线索跟进规则", 2)
    add_bullets(doc, [
        "新线索必须尽快响应，首次沟通的质量直接影响用户信任。",
        "每次跟进都要记录用户状态、核心顾虑、下一步动作和约定时间。",
        "长期未转化线索不应反复打扰，应通过内容、活动和阶段性关怀保持关系。",
        "转介绍线索要特别保护推荐人信任，跟进前先确认推荐背景和沟通边界。",
    ])
    add_heading(doc, "5.3 咨询转化底线", 2)
    add_callout(doc, "转化底线", "不夸大结果、不制造焦虑、不隐瞒关键信息、不为了成交牺牲服务可交付性。真正稳定的转化来自理解、匹配和信任。")

    add_heading(doc, "六、运营管理制度", 1)
    add_heading(doc, "6.1 校区/项目日常运营标准", 2)
    add_bullets(doc, [
        "每日检查：教室环境、设备、物资、课程安排、老师到岗、学生出勤和当天异常事项。",
        "每周检查：线索进展、在读服务、家校沟通、作业反馈、排课冲突、投诉风险和数据完整性。",
        "每月检查：收入、续费、退费、满班率、课消、师资效率、活动效果和服务满意度。",
    ])
    add_heading(doc, "6.2 合同、收款与退费流程", 2)
    add_table(
        doc,
        ["事项", "执行要求", "责任岗位", "风险提醒"],
        [
            ["合同签署", "确认课程、金额、有效期、退转规则和双方权责", "咨询/运营", "不得口头承诺与合同不一致内容"],
            ["收款确认", "缴费后及时核对金额、项目、学生姓名和入账信息", "运营/财务协同", "避免错收、漏收、项目归属不清"],
            ["转班/停课", "确认原因、剩余课时、可选方案和家长确认记录", "教务/运营", "必须保留书面或系统记录"],
            ["退费处理", "先了解原因并评估可挽回方案，再按合同规则处理", "校区负责人/运营", "情绪类退费要先安抚，避免升级投诉"],
        ],
        [1300, 3650, 1800, 2610],
    )
    add_heading(doc, "6.3 客诉处理机制", 2)
    add_steps(doc, [
        ("接住情绪", "第一时间表达重视，不争辩，不推责，先让用户感到被看见。"),
        ("还原事实", "收集合同、沟通记录、课堂记录、收费记录、教师反馈和相关人员说明。"),
        ("判断等级", "根据影响范围、情绪强度、金额风险和传播风险判断是否升级。"),
        ("给出方案", "方案应包括解释、补救、责任人、时间节点和后续跟踪方式。"),
        ("复盘沉淀", "客诉结束后复盘流程漏洞，并转化为培训案例或制度补丁。"),
    ])

    add_heading(doc, "七、人员协作机制", 1)
    add_heading(doc, "7.1 岗位责任矩阵", 2)
    add_table(
        doc,
        ["关键事项", "主责", "协同", "最终确认"],
        [
            ["线索来源与活动数据", "市场/运营", "咨询", "业务负责人"],
            ["用户诊断与方案", "咨询", "教学负责人/测评支持", "校区负责人"],
            ["排课与入学交接", "教务", "咨询/教师/班主任", "教务负责人"],
            ["课堂交付与作业反馈", "教师", "班主任/教研", "教学负责人"],
            ["家校沟通与服务维护", "班主任", "教师/咨询", "校区负责人"],
            ["合同收款与退转处理", "运营", "咨询/财务协同", "校区负责人"],
            ["AI与数据系统支持", "1605/系统支持", "各业务岗位", "技术或业务负责人"],
        ],
        [2300, 1800, 2900, 2360],
    )
    add_heading(doc, "7.2 跨部门交接标准", 2)
    add_bullets(doc, [
        "交接必须包含背景、现状、风险、承诺、下一步和截止时间。",
        "涉及用户体验的事项，不允许只口头交接；必须在系统或表单中留下记录。",
        "交接后接收人需确认是否能完成，不能完成要及时说明原因并申请支持。",
        "跨部门争议以用户利益、合同边界和组织长期价值为优先判断标准。",
    ])
    add_heading(doc, "7.3 升级机制", 2)
    add_table(
        doc,
        ["等级", "典型情况", "处理要求"],
        [
            ["一级：岗位内处理", "普通咨询、日常排课、常规学情反馈", "责任岗位当天闭环并记录"],
            ["二级：负责人介入", "家长明显不满、教师临时冲突、学生连续缺勤、方案需要调整", "24小时内给出处理方案"],
            ["三级：管理层介入", "退费争议、舆情风险、重大安全或合规风险、跨部门严重冲突", "立即升级，统一口径，形成书面记录"],
        ],
        [1500, 4600, 3260],
    )

    add_heading(doc, "八、质量管理与复盘", 1)
    add_heading(doc, "8.1 指标体系", 2)
    add_table(
        doc,
        ["指标类别", "关注内容", "示例指标"],
        [
            ["增长指标", "线索、到访、试听、报名、续费、转介绍", "有效线索数、到访率、转化率、续费率"],
            ["教学指标", "出勤、作业、测评、阶段进步、课程满意度", "出勤率、作业完成率、阶段测评变化"],
            ["服务指标", "反馈及时性、家校沟通质量、客诉响应", "反馈准时率、满意度、投诉关闭时长"],
            ["运营指标", "排课效率、课消、师资利用、数据完整性", "满班率、课消率、排课冲突数"],
            ["组织指标", "复盘质量、问题闭环、人才状态", "复盘完成率、问题关闭率、关键岗位稳定性"],
        ],
        [1700, 3900, 3760],
    )
    add_heading(doc, "8.2 复盘机制", 2)
    add_steps(doc, [
        ("周复盘", "聚焦本周线索、课程、服务、异常、数据和下周重点动作。"),
        ("月复盘", "聚焦业务结果、问题归因、流程漏洞、优秀案例和资源配置。"),
        ("项目复盘", "围绕活动、招生季、课程项目或重大客诉进行专项复盘。"),
        ("制度更新", "把复盘中重复出现的问题，沉淀为流程更新、话术模板、检查清单或培训内容。"),
    ])
    add_callout(doc, "复盘标准", "好的复盘不是开会总结情绪，而是把事实讲清、把原因挖透、把责任落下、把下一步动作排出来。")

    add_heading(doc, "九、AI与系统使用规范", 1)
    add_heading(doc, "9.1 1605人工智能的支持场景", 2)
    add_bullets(doc, [
        "教学支持：辅助生成练习、错题分析、课堂素材、阶段反馈和个性化学习建议。",
        "运营支持：辅助整理线索标签、活动复盘、用户分层和服务提醒。",
        "管理支持：辅助汇总指标、识别异常、形成复盘纪要和问题清单。",
        "内容支持：辅助生成品牌内容、课程说明、活动文案和知识库材料。",
    ])
    add_heading(doc, "9.2 数据录入规范", 2)
    add_table(
        doc,
        ["数据类型", "录入要求", "使用价值"],
        [
            ["用户基础信息", "姓名、年级、联系方式、来源、目标、关键诉求准确完整", "支持咨询跟进和用户分层"],
            ["学习过程数据", "出勤、作业、测评、课堂表现、阶段反馈及时记录", "支持学情判断和个性化服务"],
            ["沟通记录", "记录沟通对象、时间、问题、结论和下一步", "支持服务连续性和风险追溯"],
            ["异常与客诉", "记录事实、责任人、处理过程和关闭结果", "支持风险管理和流程改进"],
        ],
        [1800, 4550, 3010],
    )
    add_heading(doc, "9.3 AI使用边界", 2)
    add_bullets(doc, [
        "AI可以辅助分析和生成，但不能替代教师、咨询和管理者的专业判断。",
        "涉及学生隐私、家庭信息、合同金额和敏感沟通内容时，必须遵守内部权限和数据安全要求。",
        "AI生成内容对外使用前必须人工审核，尤其是学习建议、宣传内容和家校反馈。",
        "系统数据的价值取决于一线录入质量，缺失或失真的数据会影响后续判断。",
    ])

    add_heading(doc, "十、附录工具包", 1)
    add_heading(doc, "10.1 建议配套表单", 2)
    add_table(
        doc,
        ["表单名称", "使用场景", "核心字段"],
        [
            ["用户咨询信息表", "首次咨询和线索建档", "来源、学生信息、目标、痛点、家长诉求、下一步"],
            ["学习诊断记录表", "测评后或试听后", "基础水平、问题诊断、风险点、推荐方案"],
            ["入学交接单", "报名后交给教务/教学/班主任", "承诺事项、学生画像、排课安排、沟通偏好"],
            ["课堂反馈表", "课后记录", "出勤、掌握情况、作业、异常、需跟进事项"],
            ["阶段学情报告", "阶段复盘", "目标、进展、证据、问题、调整方案"],
            ["客诉处理记录表", "投诉或重大不满", "事实、诉求、责任、方案、关闭结果"],
            ["活动复盘表", "市场活动结束后", "目标、投入、到访、留资、转化、问题、下次优化"],
            ["月度运营复盘表", "月度经营会议", "指标、亮点、问题、归因、动作、责任人"],
        ],
        [2200, 2700, 4460],
    )
    add_heading(doc, "10.2 模板示例：入学交接单", 2)
    add_table(
        doc,
        ["字段", "填写说明"],
        [
            ["学生基本信息", "姓名、年级、学校、目标、当前基础"],
            ["家长核心诉求", "最关心的问题、期望结果、沟通偏好"],
            ["学生状态画像", "学习习惯、情绪状态、自驱力、优势与短板"],
            ["推荐方案", "课程/班型/老师/服务节奏/阶段目标"],
            ["已承诺事项", "咨询过程中明确承诺过的内容，必须如实交接"],
            ["风险提示", "家长敏感点、学生抗拒点、时间冲突、费用或结果预期风险"],
            ["首周重点", "首次上课提醒、课后反馈时间、需要重点观察的事项"],
        ],
        [2300, 7060],
    )
    add_heading(doc, "10.3 模板示例：周复盘问题清单", 2)
    add_bullets(doc, [
        "本周新增线索中，哪些来源质量最高？原因是什么？",
        "本周咨询未转化的主要卡点是什么？是否是产品、价格、信任、时间或方案问题？",
        "本周教学服务中，哪些学生出现异常？是否已经有人负责跟进？",
        "本周家长反馈中，有没有重复出现的问题？是否需要调整流程？",
        "本周有哪些值得复制的优秀案例？是否已经沉淀为话术、模板或培训材料？",
        "下周最重要的三件事是什么？责任人和截止时间分别是什么？",
    ])
    add_heading(doc, "结语", 1)
    add_p(doc, "流程和制度的目的，不是把人变成机械执行者，而是让组织在复杂场景中仍然保持稳定、清晰和可信。真正有效的制度，最终会变成一线同事做对事情的默认路径，也会成为用户感受到专业、温度和确定性的来源。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
