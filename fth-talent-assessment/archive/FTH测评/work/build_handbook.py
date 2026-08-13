from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/chenpan/Documents/于无声处起惊雷")
OUTPUT = ROOT / "output" / "于无声处起惊雷-小凡教育科技员工手册-VI品牌版-V1.2.docx"
LOGO = ROOT / "work" / "assets" / "feifan-logo-black.png"

FONT = "HarmonyOS Sans SC"
TITLE_FONT = "方正清刻本悦宋简体"
CREAM = "FFFCE9"
PAPER = "FFFFFF"
INK = "1D1D1D"
MUTED = "777777"
FOREST = "2E3192"
MOSS = "55B6FF"
YELLOW = "FFE200"
PALE_YELLOW = "FFF6A8"
PALE_GREEN = "EAF6FF"
WHITE = "FFFFFF"
LINE = "D7D7C8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_width(table, widths_cm):
    table.autofit = False
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "140")
    tbl_ind.set(qn("w:type"), "dxa")


def set_run(run, size=None, color=INK, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def shade_paragraph(paragraph, fill, color=INK, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), border)
        p_bdr.append(bottom)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_text(doc, text="", size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=7, line=1.35,
             keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_kicker(doc, text):
    p = add_text(doc, text.upper(), size=8.5, color=MOSS, bold=True, after=4, keep=True)
    p.paragraph_format.left_indent = Cm(0.05)
    return p


def add_h1(doc, text, kicker=None):
    if kicker:
        add_kicker(doc, kicker)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), font=TITLE_FONT)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(text), size=10.3)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(text), size=10.3)
    return p


def add_callout(doc, label, text, fill=PALE_YELLOW, accent=YELLOW):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [0.35, 17.15])
    left, right = table.rows[0].cells
    set_cell_shading(left, accent)
    set_cell_shading(right, fill)
    for cell in (left, right):
        set_cell_margins(cell, top=170, bottom=170, start=160, end=160)
        set_cell_border(
            cell,
            top={"val": "nil"}, bottom={"val": "nil"},
            left={"val": "nil"}, right={"val": "nil"},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(label + "  "), size=9.5, bold=True, color=FOREST)
    set_run(p.add_run(text), size=10.2, color=INK)
    add_text(doc, "", size=1, after=3)
    return table


def add_behavior_block(doc, title, meaning, actions, boundary):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [3.25, 14.25])
    left, right = table.rows[0].cells
    set_cell_shading(left, FOREST)
    set_cell_shading(right, PAPER)
    for cell in (left, right):
        set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "5", "color": LINE},
            bottom={"val": "single", "sz": "5", "color": LINE},
            left={"val": "single", "sz": "5", "color": LINE},
            right={"val": "single", "sz": "5", "color": LINE},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_after = Pt(0)
    set_run(lp.add_run(title), size=15, color=YELLOW, bold=True)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(6)
    set_run(rp.add_run(meaning), size=10.2, color=INK, bold=True)
    for item in actions:
        p = right.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(item), size=9.7, color=INK)
    bp = right.add_paragraph()
    bp.paragraph_format.space_before = Pt(5)
    bp.paragraph_format.space_after = Pt(0)
    set_run(bp.add_run("边界："), size=9.5, bold=True, color=MOSS)
    set_run(bp.add_run(boundary), size=9.5, color=MUTED)
    add_text(doc, "", size=1, after=4)


def add_section_page(doc, number, kicker, title, statement):
    doc.add_page_break()
    add_text(doc, f"{number:02d}", size=48, color=YELLOW, bold=True, after=3)
    add_kicker(doc, kicker)
    p = add_text(doc, title, size=27, color=FOREST, bold=True, after=12, line=1.05)
    for run in p.runs:
        set_run(run, size=27, color=FOREST, bold=True, font=TITLE_FONT)
    keep_with_next(p)
    add_text(doc, statement, size=13, color=INK, bold=True, after=18, line=1.35)
    p = add_text(doc, "FOREST FOR TALENTS", size=8.5, color=MOSS, bold=True, after=30)
    shade_paragraph(p, CREAM)


def add_three_column(doc, items):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [5.83, 5.83, 5.84])
    for idx, (head, sub, body) in enumerate(items):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_GREEN if idx != 1 else PALE_YELLOW)
        set_cell_margins(cell, top=220, bottom=220, start=200, end=200)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": WHITE},
            bottom={"val": "single", "sz": "4", "color": WHITE},
            left={"val": "single", "sz": "4", "color": WHITE},
            right={"val": "single", "sz": "4", "color": WHITE},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(head), size=13, bold=True, color=FOREST)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        set_run(p2.add_run(sub), size=8.5, color=MOSS, bold=True)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        p3.paragraph_format.line_spacing = 1.3
        set_run(p3.add_run(body), size=9.5, color=INK)
    add_text(doc, "", size=1, after=4)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("小凡教育科技  |  于无声处起惊雷  |  VI品牌版 V1.2  |  "), size=8, color=MUTED)
    run = p.add_run()
    set_run(run, size=8, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.2))
    else:
        set_run(p.add_run("非凡教育"), size=7.5, color=INK, bold=True)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.35

    h1 = styles["Heading 1"]
    h1.font.name = TITLE_FONT
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_FONT)
    h1.font.size = Pt(21)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(FOREST)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = TITLE_FONT
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_FONT)
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(FOREST)
    h2.paragraph_format.space_before = Pt(11)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.3)
        style.font.color.rgb = RGBColor.from_string(INK)


def build():
    doc = Document()
    doc.core_properties.title = "于无声处起惊雷：小凡教育科技员工手册"
    doc.core_properties.subject = "小凡教育科技文化、组织与协作手册"
    doc.core_properties.author = "小凡教育科技"
    doc.core_properties.keywords = "员工手册, 青色组织, 人才森林计划, FTH职业特质"
    doc.core_properties.comments = "VI品牌版 V1.2"
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    add_header(section)
    add_footer(section)
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    # Cover
    logo_p = doc.add_paragraph()
    logo_p.paragraph_format.space_after = Pt(58)
    if LOGO.exists():
        logo_p.add_run().add_picture(str(LOGO), width=Inches(2.25))
    else:
        set_run(logo_p.add_run("非凡教育"), size=11, color=INK, bold=True)
    add_kicker(doc, "EMPLOYEE HANDBOOK")
    cover_title = add_text(doc, "于无声处\n起惊雷", size=38, color=INK, bold=True, after=16, line=0.98)
    for run in cover_title.runs:
        set_run(run, size=38, color=INK, bold=True, font=TITLE_FONT)
    shade_paragraph(cover_title, YELLOW)
    add_text(doc, "员工手册", size=17, color=INK, bold=True, after=8)
    add_text(doc, "VI品牌版  V1.2", size=10.5, color=FOREST, bold=True, after=62)
    p = add_text(
        doc,
        "我们想共同创造的，不只是一家公司，\n而是一片让人扎根、生长、开花的森林。",
        size=13, color=INK, bold=True, line=1.45, after=25,
    )
    shade_paragraph(p, CREAM, border=YELLOW)
    add_text(doc, "90分 Native  +  AI Native  =  未来组织", size=10.5, color=FOREST, bold=True, after=4)
    add_text(doc, "2026 · 内部使用", size=8.5, color=MUTED, after=0)

    # Front matter
    doc.add_page_break()
    add_h1(doc, "写在前面", "A LETTER TO EVERY COLLEAGUE")
    add_text(
        doc,
        "欢迎来到非凡。你手里的不是一本要求你“照做”的说明书，而是一份我们共同工作的约定。"
        "它试图回答三个问题：我们为什么在一起？我们怎样一起把事情做好？我们如何让个人与组织一起成长？",
        size=11, after=10,
    )
    add_callout(
        doc,
        "我们相信",
        "真正有生命力的组织，不靠层层命令推动，而靠共同底色、清晰目标、真实信息和愿意负责的人持续进化。",
    )
    add_text(
        doc,
        "小凡教育科技深耕教育行业 18 年，是一家以认知科学为内核、AI 科技为驱动的全链路教育科技集团。"
        "非凡教育是集团成熟业务标杆，长期陪伴一群常被传统节奏忽略的学生，帮助他们提分、走进理想大学、"
        "重新被尊重，也突破自己曾经以为的极限。今天，AI 正在重写教育与工作，我们要建设一支既有 90 分"
        "品味与能力、又能用 AI 完成端到端创造的人才组织。",
    )
    add_text(
        doc,
        "这本手册记录的是现阶段已经形成共识的文化、人才与协作方式。它会随着业务、团队和我们对教育的理解继续生长。"
        "当本手册与劳动合同、公司正式制度或法律法规不一致时，以后者为准。",
        size=9.3, color=MUTED, italic=True,
    )

    add_h2(doc, "怎样使用这本手册")
    add_bullet(doc, "入职时，用它理解非凡的底色、业务和工作方式。")
    add_bullet(doc, "遇到分歧时，用它回到共同语言，而不是猜测彼此动机。")
    add_bullet(doc, "做重要决定时，用它检查：信息是否真实、建议是否充分、责任是否清楚。")
    add_bullet(doc, "每次复盘时，用它追问：我为客户、团队和长期组织多存入了什么价值？")

    doc.add_page_break()
    add_h1(doc, "一张图读懂非凡", "THE HANDBOOK MAP")
    add_three_column(doc, [
        ("底色", "BELIEFS", "决定我们能否成为长期战友。共同底色是善良、阳光、自驱、自燃，以及诚实、自省、担当、酷。"),
        ("特质", "WAYS OF CONTRIBUTING", "决定我们更自然的贡献方式。Fighter 推进结果，Thinker 构建认知，Helper 赋能人和系统。"),
        ("技能", "CAPABILITIES", "决定我们能创造什么价值。技能需要在真实项目、用户反馈、复盘和 AI 协作中持续升级。"),
    ])
    add_callout(doc, "选择顺序", "先看底色是否同频，再看特质是否适配，最后用训练与实战把技能练到 90 分。", fill=PALE_GREEN, accent=FOREST)
    add_h2(doc, "手册的三条主线")
    add_number(doc, "为什么在一起：理解时代、使命、用户与共同价值观。")
    add_number(doc, "怎样一起做事：理解青色组织、自主管理、建议流程与协作边界。")
    add_number(doc, "如何一起成长：理解人才画像、FTH 特质、人才森林计划与学习方法。")

    # 01
    add_section_page(doc, 1, "WHY NOW", "我们身处的时代", "复制和低价竞争正在失效。未来属于能定义高标准、创造真实价值并借助 AI 完成闭环的人。")
    add_h1(doc, "从增量时代走向存量时代")
    add_text(
        doc,
        "过去，市场增长常常能掩盖产品、服务与组织能力的不足。今天，用户拥有更多选择，信息更加透明，"
        "“差不多”越来越难以赢得信任。非凡要做的不是更忙，而是更准确地理解用户、更有品味地定义目标、"
        "更扎实地把价值交付到最后一公里。",
    )
    add_h2(doc, "90分 Native")
    add_text(
        doc,
        "能识别什么是真正的好，能把模糊问题定义清楚，能在约束中完成高质量交付。90 分不是完美主义，"
        "而是对用户价值、判断质量和执行细节保持高标准。",
    )
    add_h2(doc, "AI Native")
    add_text(
        doc,
        "不把 AI 当作偶尔使用的搜索框，而把它当作新的工作基础设施：主动学习、设计工作流、构建智能体，"
        "从理念、内容、产品到用户反馈形成端到端闭环。AI 放大能力，也会放大判断；因此，品味、责任与事实核验比过去更重要。",
    )
    add_callout(doc, "共同命题", "用更少的无效消耗，创造更高质量的教育体验；用技术扩大个性化，用人保留温度。")

    # 02
    add_section_page(doc, 2, "WHO WE ARE", "我们是谁", "我们属于小凡教育科技：用认知科学理解学习，用 AI 放大教育能力，用真实教育实践持续验证。")
    add_h1(doc, "探索教育更多可能性")
    add_text(
        doc,
        "小凡教育科技是一家以认知科学为内核、AI 科技为驱动的全链路教育科技集团。深耕教育行业 18 年，"
        "我们以“探索教育更多可能性”为核心愿景，构建“科技 + 科学 + 教育”三位一体的核心壁垒，"
        "打通从认知科学研究、AI 科技研发到真实教育实践的完整闭环。",
    )
    add_callout(
        doc,
        "我们的独特性",
        "科研能够落地，技术拥有真实场景，教育模式可以被持续验证。我们既不是传统教培机构的简单升级，"
        "也不是脱离教育现场的纯技术研发公司。",
        fill=PALE_GREEN,
        accent=FOREST,
    )
    add_h2(doc, "集团业务版图")
    group = doc.add_table(rows=5, cols=3)
    group.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(group, [3.5, 4.2, 9.8])
    group_rows = [
        ("载体", "核心角色", "正在形成的价值"),
        ("1605 研发中心", "技术核心底座", "建设百万级学生全周期认知行为专属数据集，沉淀 AI 技术、教育产品和自主知识产权，并与国内 985 高校脑科学实验室共建产学研体系。"),
        ("非凡教育", "成熟业务标杆", "深耕浙江高考提分赛道，把教学、陪伴与科技结合，持续打造可量化、可复制的精准教育模型。"),
        ("小凡私塾", "未来教育探索载体", "汇聚不同领域的专家工作室，探索心流学习、认知塑造等面向人的完整成长议题。"),
        ("教育 SaaS", "行业赋能平台", "把在真实教育现场验证过的产品、方法和技术能力，转化为面向行业伙伴的教育科技支持。"),
    ]
    for r_idx, row in enumerate(group_rows):
        for c_idx, text in enumerate(row):
            cell = group.rows[r_idx].cells[c_idx]
            set_cell_margins(cell, top=130, bottom=130, start=120, end=120)
            set_cell_shading(cell, FOREST if r_idx == 0 else (PALE_GREEN if r_idx % 2 else PAPER))
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.22
            set_run(
                p.add_run(text),
                size=8.8,
                color=WHITE if r_idx == 0 else INK,
                bold=(r_idx == 0 or c_idx < 2),
            )
    set_repeat_table_header(group.rows[0])
    add_text(doc, "", size=1, after=3)
    add_h2(doc, "从数据到教育现场")
    add_text(
        doc,
        "集团以产业、学校与高校、科研机构三方联动的产学研方式，让研究成果进入业务，并在真实学习过程里"
        "获得反馈。阶段性数据表明：截至 2026 年 4 月 3 日，本学年学练机累计收集学习数据时长已达 "
        "1,252,518 分钟；相关 AI 应用日均消耗 Tokens 约 1.5 亿。数据不是目的，它帮助我们更准确地理解"
        "学习行为、验证教育方法，并持续改善产品与服务。",
    )
    add_text(
        doc,
        "注：上述数据为特定统计口径下的阶段数据，后续以集团最新正式披露为准。自主知识产权包括 AI 技术代码、"
        "自编教材等相关成果。",
        size=8.8,
        color=MUTED,
        italic=True,
    )
    add_h1(doc, "让沉默发声，让平凡不凡")
    add_text(
        doc,
        "非凡教育始于艺考生文化课冲刺。我们看见这些情感丰富、渴望被理解的学生，在传统学校的统一节奏中容易掉队。"
        "于是，我们用更适合他们的时间、节奏、教学和陪伴，帮助他们重新进入学习、重新获得尊重。",
    )
    add_h2(doc, "我们的教学理念")
    add_three_column(doc, [
        ("陪你", "STAY WITH YOU", "不轻易放弃任何一个学生。先建立连接，再推动改变。"),
        ("教你", "TEACH YOU", "不仅教知识，也教面对情绪、拆解困难和持续行动。"),
        ("逼你 · 鼓励你", "PUSH & ENCOURAGE", "温和而坚定地推动行动，让每一次努力被看见、被反馈。"),
    ])
    add_h2(doc, "我们的价值交付")
    add_bullet(doc, "提分：让努力转化为可见的学习结果。")
    add_bullet(doc, "考上更好的大学：让选择空间真正扩大。")
    add_bullet(doc, "重新被尊重：让学生看见自己的能力与可能。")
    add_bullet(doc, "突破极限：把一次备考变成面对未来的成长经验。")
    add_h2(doc, "非凡教育与集团的共同进化")
    add_text(
        doc,
        "非凡教育既是集团成熟业务，也是科学与科技进入真实教育场景的重要验证场。我们持续建设 AI 学练与"
        "用户运营能力，让“千人千面、实时反馈、游戏化激励”从理念变成系统，并把验证有效的方法沉淀为"
        "集团产品与行业能力。业务可以扩张，但我们不能丢掉最初的东西：看见具体的人。",
    )

    # 03
    add_section_page(doc, 3, "TEAL ORGANIZATION", "我们想成为怎样的组织", "把组织当作一片会生长的森林：有共同土壤，也允许每棵树用自己的方式生长。")
    add_h1(doc, "青色组织不是“没有管理”")
    add_text(
        doc,
        "青色组织的核心不是取消规则，更不是每个人只做自己喜欢的事。它改变的是权力与责任的关系："
        "让信息更透明、让决定更靠近问题、让人带着完整的自己工作，并让组织持续倾听真正的使命。",
    )
    add_three_column(doc, [
        ("自主管理", "SELF-MANAGEMENT", "决定尽量由最接近问题、最愿意承担结果的人发起；自由必须与建议、透明和闭环同时出现。"),
        ("身心完整", "WHOLENESS", "不必戴着完美的职业面具。可以表达事实、感受与不确定，也要尊重边界并对自己的行为负责。"),
        ("进化宗旨", "EVOLUTIONARY PURPOSE", "战略不是只靠少数人预测未来，而是持续倾听用户、产品、数据和组织正在召唤我们成为什么。"),
    ])
    add_h2(doc, "草原与狮子")
    add_text(
        doc,
        "公司努力提供公平的土壤、清晰的规则、必要的资源与相对自由的空间，这就是“草原”。"
        "每个人要成为主动寻找目标、敢于承担结果、能够自主协作的“狮子”。只有草原没有狮子，组织会失去结果；"
        "只有狮子没有草原，团队会陷入争抢和内耗。",
    )
    add_callout(doc, "自由的定义", "自由不是少做事，而是拥有选择方法的空间，同时对承诺、影响和结果承担更多责任。")

    # 04
    add_section_page(doc, 4, "HOW WE WORK", "我们如何一起做事", "少一点等待指令，多一点看见问题、发起行动、征求建议和完成闭环。")
    add_h1(doc, "决定要靠近问题")
    add_text(
        doc,
        "最了解问题的人，通常最有机会做出高质量判断。非凡鼓励每个人成为发起者，但发起不等于独断。"
        "越重要、越不可逆、影响越广的决定，越需要更充分的建议与正式治理。",
    )
    add_h2(doc, "三类决定")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [3.1, 6.0, 8.4])
    rows = [
        ("类型", "适用情形", "推荐做法"),
        ("A｜可逆、局部", "影响小、成本低、容易撤回", "直接决定并同步相关人；用结果验证。"),
        ("B｜跨团队、较重要", "影响他人资源、目标、体验或流程", "发起建议流程；记录依据、建议与最终决定。"),
        ("C｜重大、不可逆", "涉及法律、安全、品牌、重大预算或公司级方向", "按公司正式授权与审批机制执行，并充分征求专业意见。"),
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_margins(cell)
            set_cell_shading(cell, FOREST if r_idx == 0 else PAPER)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": LINE},
                bottom={"val": "single", "sz": "5", "color": LINE},
                left={"val": "single", "sz": "5", "color": LINE},
                right={"val": "single", "sz": "5", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(text), size=9.3, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0 or c_idx == 0))
    set_repeat_table_header(table.rows[0])

    add_h2(doc, "建议流程：任何人都可以发起，但不能跳过倾听")
    add_number(doc, "定义问题：你看见了什么事实？希望解决什么，而不是先证明谁错了。")
    add_number(doc, "识别影响：谁会受到影响？谁拥有专业经验、关键数据或必要授权？")
    add_number(doc, "征求建议：清楚说明背景、约束、选项与截止时间，认真听取不同意见。")
    add_number(doc, "作出决定：发起者综合建议后决定；不需要所有人同意，但需要能说明理由。")
    add_number(doc, "公开闭环：同步决定、负责人、时间点和复盘方式，并承担后果。")
    add_callout(doc, "建议不是投票", "被征求建议的人拥有表达权，不自动拥有否决权；发起者拥有决定权，也承担结果责任。", fill=PALE_GREEN, accent=FOREST)

    add_h2(doc, "信息透明：坏消息要跑在问题前面")
    add_bullet(doc, "事实、判断和情绪分开表达。先说我们知道什么，再说我们如何理解。")
    add_bullet(doc, "遇到风险及时暴露，不包装成“进展顺利”，也不等到无法挽回才升级。")
    add_bullet(doc, "共享信息以完成工作为目的，同时遵守隐私、保密和最小必要原则。")
    add_bullet(doc, "公开承诺与进度，让协作方不需要反复追问。")

    add_h2(doc, "会议要产生下一步")
    add_number(doc, "会前说明目的：同步、讨论、决策或复盘，不把四种目的混在一起。")
    add_number(doc, "会中先事实后观点，允许短暂的不同意见，不用职位结束讨论。")
    add_number(doc, "会后明确行动：谁、在什么时候、交付什么、如何判断完成。")
    add_number(doc, "能异步完成的，不开会；需要真实碰撞的，不躲在消息里。")

    add_h2(doc, "冲突处理：直接、尊重、逐级获得支持")
    add_number(doc, "先和当事人直接沟通，描述具体行为、影响和你的请求。")
    add_number(doc, "若无法推进，邀请双方都信任的同事或 HRBP 作为调解者。")
    add_number(doc, "若仍无法解决，邀请与事件无直接利益冲突的伙伴组成评议小组。")
    add_number(doc, "涉及违法、骚扰、安全、舞弊等事项，立即走正式举报与处理机制，不要求员工自行调解。")

    # 05
    add_section_page(doc, 5, "OUR VALUES", "我们的共同底色", "价值观不是墙上的四个词，而是压力来临时，我们仍愿意选择的行为。")
    add_h1(doc, "诚实 · 自省 · 担当 · 酷")
    add_behavior_block(
        doc, "诚实",
        "尽可能准确、完整、及时地提供信息，让团队能够基于现实协作。",
        ["说清事实、判断、风险和不知道的部分。", "发现承诺可能无法完成时，提前沟通并提出新方案。", "面对用户与伙伴，不用模糊表达掩盖问题。"],
        "诚实不等于情绪化地“有话直说”，更不等于用真实之名伤害他人。",
    )
    add_behavior_block(
        doc, "自省",
        "能跳出当下感受，看见自己的模式、影响以及与目标之间的差距。",
        ["先问“我能改变什么”，再讨论环境和他人的责任。", "用数据、反馈和复盘修正自我判断。", "承认不知道、做错了或需要帮助。"],
        "自省不是自我否定，也不是把所有问题都归咎于自己。",
    )
    add_behavior_block(
        doc, "担当",
        "敢于承诺，也能组织资源、管理风险并把事情真正闭环。",
        ["承诺前理解目标和边界，承诺后持续同步进度。", "遇到阻力主动求助、拆解路径，而不是安静地失约。", "对结果负责，也对过程中影响到的人负责。"],
        "担当不是一个人硬扛。长期透支、隐瞒风险和拒绝协作都不是担当。",
    )
    add_behavior_block(
        doc, "酷",
        "热爱且忠诚，偏好挑战，创造性执行，并在困难里保持建设性的乐观。",
        ["不满足于复制旧答案，敢于提出更好的可能。", "选择难而正确的事，并把创意落到用户体验。", "胜利时分享功劳，受挫时保持行动。"],
        "酷不是特立独行的姿态，也不是为了新奇而忽视结果、规则和他人。",
    )
    add_h2(doc, "存取：先创造价值，再期待回报")
    add_text(
        doc,
        "我们把信任、能力、影响力和长期合作想象成一只共同账户。每一次可靠交付、主动补位、真诚反馈和知识沉淀，"
        "都是存入；每一次占用资源、寻求支持和获得机会，都是取出。健康的关系不是只存不取，而是长期保持正向循环。",
    )

    # 06
    add_section_page(doc, 6, "TALENT PROFILE", "我们寻找什么样的人", "先看底色，再看特质，最后持续升级技能。人不是标签，而是会生长的生命。")
    add_h1(doc, "A+ 人才：善良、阳光、自驱、自燃")
    add_text(
        doc,
        "我们欣赏真正打赢过胜仗的人。这里的“胜仗”不只是一张漂亮履历，而是你在真实限制中完成过结果、"
        "承担过后果、形成过方法。这样的自信更安静，也更可靠。",
    )
    add_h2(doc, "自驱力的四个发动机")
    add_three_column(doc, [
        ("热爱", "ARDENT LOVE", "我对这件事有真实兴趣，愿意持续投入，而不只依赖外部奖励。"),
        ("胜任", "COMPETENCE", "我具备完成任务的能力，也愿意承认差距并训练。"),
        ("责任 · 归属", "RESPONSIBILITY & BELONGING", "我愿意承担影响，也相信自己是共同体的一部分。"),
    ])
    add_callout(doc, "自燃状态", "当热爱、胜任、责任与归属同时点亮，人会从“被要求完成”进入“被梦想叫醒，以责任前行”。")

    add_h1(doc, "FTH：六种贡献方式")
    add_text(
        doc,
        "FTH 测评用于发展对话、岗位建议和团队搭配，不用于给人贴标签。主分型代表更自然的发力方式，"
        "次分型代表复合能力，第三倾向可作为未来培养方向。",
    )
    fth = doc.add_table(rows=7, cols=4)
    fth.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(fth, [2.0, 3.0, 5.1, 7.4])
    fth_rows = [
        ("特质", "分型", "自然贡献", "典型场景"),
        ("F 进取者", "冲刺型 Runner", "快速行动、快速反馈、把机会变成结果", "招生冲刺、市场活动、社群增长、短期项目"),
        ("F 进取者", "攻坚型 Climber", "在阻力与压力中持续推进、拿下硬仗", "重点用户、复杂沟通、危机处理、关键项目"),
        ("T 思辨者", "分析型 Analyzer", "从复杂信息中找到规律、根因与优先级", "用户与学情分析、经营复盘、课程效果评估"),
        ("T 思辨者", "创构型 Builder", "把想法搭成课程、系统、机制与方法", "教研体系、运营机制、知识库、SOP、组织设计"),
        ("H 赋能者", "人际型 Socializer", "理解人、连接人、激发人并恢复合作", "教学与教练、家校沟通、HRBP、跨部门协调"),
        ("H 赋能者", "流程型 Keeper", "维护秩序、细节与长期稳定交付", "教务排课、流程运营、质量管理、数据与档案"),
    ]
    for r_idx, row in enumerate(fth_rows):
        for c_idx, text in enumerate(row):
            cell = fth.rows[r_idx].cells[c_idx]
            set_cell_margins(cell, top=110, bottom=110, start=110, end=110)
            set_cell_shading(cell, FOREST if r_idx == 0 else (PALE_GREEN if r_idx % 2 else PAPER))
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            set_run(p.add_run(text), size=8.6, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0 or c_idx < 2))
    set_repeat_table_header(fth.rows[0])
    add_h2(doc, "好团队不是同类相加")
    add_text(
        doc,
        "Fighter 让目标进入行动，Thinker 让复杂问题变清楚，Helper 让人和系统长期运转。"
        "重要项目应主动检查：是否有人推进结果、有人判断方向、有人托住协作与交付？",
    )

    # 07
    add_section_page(doc, 7, "GROWTH SYSTEM", "我们如何一起成长", "成长不是等培训发生，而是选择方向、进入实战、获得反馈、形成方法。")
    add_h1(doc, "人才森林计划")
    add_three_column(doc, [
        ("15天验证期", "SEED", "理解底色与业务，进入真实现场，验证彼此是否适合长期同行。"),
        ("3—6个月成长期", "SPROUT", "明确岗位结果，完成关键任务，在反馈中形成稳定打法。"),
        ("长期发展期", "TREE", "从完成任务走向创造系统，从个人贡献走向带动他人和组织进化。"),
    ])
    add_h2(doc, "成长的四个来源")
    add_bullet(doc, "一带一：由更有经验的伙伴提供现场反馈与方法示范。")
    add_bullet(doc, "真实项目：在用户、产品和经营问题中训练，而不是只听知识。")
    add_bullet(doc, "轮岗体验：理解业务链路，找到特质与岗位更好的结合点。")
    add_bullet(doc, "复盘文化：把一次经验沉淀为下一次可以复用的方法。")

    add_h1(doc, "西蒙学习法：动机 × 方法 × 时间")
    add_text(
        doc,
        "学习效果由三个因素相乘：积极的学习动机、有效的学习方法、必要的时间投入。"
        "其中任何一项接近零，成长都会停住。",
    )
    add_h2(doc, "四步学习闭环")
    add_number(doc, "做选择：选定真正值得投入、与岗位和使命相关的主题。")
    add_number(doc, "设目标：明确学到什么、做到什么、何时完成、如何验证。")
    add_number(doc, "会拆分：把大目标拆成小模块、小任务和短反馈周期。")
    add_number(doc, "能集中：为关键练习留出不被打扰的时间块，用作品而不是“学过”证明成长。")
    add_callout(doc, "AI 学习原则", "让 AI 帮你加速搜索、练习、模拟与反馈，但关键判断、事实核验和最终责任不能外包。", fill=PALE_GREEN, accent=FOREST)

    # 08
    add_section_page(doc, 8, "PERFORMANCE", "我们如何看待结果", "结果重要，产生结果的方式同样重要；短期胜利不能以透支用户、伙伴和组织为代价。")
    add_h1(doc, "四个维度看贡献")
    add_three_column(doc, [
        ("业务结果", "RESULT", "是否完成对用户与经营真正重要的结果，而不是只完成动作数量。"),
        ("过程质量", "QUALITY", "判断是否可靠、协作是否透明、交付是否可持续、风险是否被管理。"),
        ("成长与共建", "GROWTH & TEAM", "是否持续升级能力、沉淀方法、帮助伙伴，并让组织下一次做得更好。"),
    ])
    add_h2(doc, "绩效讨论应该发生在日常")
    add_bullet(doc, "目标开始前：确认结果、边界、资源与衡量方式。")
    add_bullet(doc, "执行过程中：及时同步进度、风险和需要的支持。")
    add_bullet(doc, "阶段结束后：讨论事实、方法、学习与下一步，不把反馈积压到年末。")
    add_bullet(doc, "出现差距时：先共同诊断是目标、能力、资源、协作还是意愿问题，再决定行动。")
    add_text(
        doc,
        "薪酬、奖金、晋升、调岗及绩效结果，以公司当期正式制度、岗位方案和书面通知为准。"
        "本手册只说明我们希望坚持的判断原则，不替代具体政策。",
        size=9.3, color=MUTED, italic=True,
    )

    # 09
    add_section_page(doc, 9, "FIRST 90 DAYS", "在非凡的前90天", "先理解，再贡献；先进入真实现场，再形成自己的打法。")
    add_h1(doc, "一段双向验证的旅程")
    timeline = doc.add_table(rows=4, cols=3)
    timeline.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(timeline, [3.0, 5.3, 9.2])
    timeline_rows = [
        ("阶段", "你要完成什么", "建议关注"),
        ("第 1—15 天", "理解与验证", "认识用户、业务链路、共同底色和岗位期待；完成首次现场任务与反馈。"),
        ("第 16—45 天", "独立交付", "承担清晰结果；使用建议流程；建立稳定的同步、复盘和求助方式。"),
        ("第 46—90 天", "形成打法", "完成一个可复盘的关键成果；明确优势、短板、FTH 发力方式与下一阶段成长目标。"),
    ]
    for r_idx, row in enumerate(timeline_rows):
        for c_idx, text in enumerate(row):
            cell = timeline.rows[r_idx].cells[c_idx]
            set_cell_margins(cell, top=150, bottom=150)
            set_cell_shading(cell, FOREST if r_idx == 0 else PAPER)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": LINE},
                bottom={"val": "single", "sz": "5", "color": LINE},
                left={"val": "single", "sz": "5", "color": LINE},
                right={"val": "single", "sz": "5", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            set_run(p.add_run(text), size=9.3, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0 or c_idx == 0))
    set_repeat_table_header(timeline.rows[0])

    add_h2(doc, "新伙伴的五个主动")
    add_bullet(doc, "主动见用户：不要只在内部材料里认识业务。")
    add_bullet(doc, "主动问清目标：不在模糊承诺中开始重要工作。")
    add_bullet(doc, "主动暴露不知道：越早求助，学习成本越低。")
    add_bullet(doc, "主动要反馈：询问哪些做得好、哪里影响了结果、下一次如何更好。")
    add_bullet(doc, "主动沉淀：把个人经验变成团队可复用的知识。")

    # 10
    add_section_page(doc, 10, "MUTUAL COMMITMENT", "我们彼此的承诺", "组织承诺提供土壤，员工承诺主动生长；信任只有在双方都持续兑现时才成立。")
    add_h1(doc, "组织向员工承诺")
    add_bullet(doc, "尽力提供公平的机会、清晰的规则和与岗位相匹配的资源。")
    add_bullet(doc, "尊重每个人的独特性，不用单一模板定义所有人的价值。")
    add_bullet(doc, "让重要信息尽可能靠近需要它的人，减少无意义的信息壁垒。")
    add_bullet(doc, "允许提出不同意见、承认错误和寻求帮助，不因真实表达而羞辱他人。")
    add_bullet(doc, "在业务允许的范围内，为成长、轮岗、实战和长期发展创造机会。")
    add_bullet(doc, "对违法、骚扰、舞弊、安全与严重失信行为建立正式边界并及时处理。")

    add_h1(doc, "员工向组织与伙伴承诺")
    add_bullet(doc, "把用户价值放在动作数量之前，把长期信任放在短期便利之前。")
    add_bullet(doc, "在拥有自由时承担责任，在需要资源时清楚求助。")
    add_bullet(doc, "准确、及时地共享信息，不让伙伴在错误信息上继续工作。")
    add_bullet(doc, "尊重差异，直接解决问题，不通过传播情绪建立同盟。")
    add_bullet(doc, "持续学习并善用 AI，提高个人产出，也让团队能力得到沉淀。")
    add_bullet(doc, "保护学生、家长、同事与公司的隐私、数据、品牌和合法权益。")
    add_callout(doc, "共同标准", "集体成功时分享荣誉；伙伴陷入困难时不旁观；出现错误时先止损、再复盘、再让系统变得更好。")

    # Closing
    doc.add_page_break()
    add_text(doc, "JOIN US", size=10, color=MOSS, bold=True, after=42)
    add_text(doc, "于无声处\n起惊雷", size=36, color=FOREST, bold=True, after=18, line=1.0)
    add_text(
        doc,
        "真正的惊雷，不一定从喧闹处开始。\n"
        "它可能从一次诚实的反馈、一个主动承担的决定、\n"
        "一次对学生不放弃的陪伴、一个被认真做完的细节开始。",
        size=13, color=INK, bold=True, after=25, line=1.5,
    )
    add_text(
        doc,
        "我们或许不是规模最大的公司，\n"
        "但我们希望成为最有温度、最有冲劲、最像伙伴的团队之一。\n\n"
        "愿你在这里扎根、生长、开花，\n"
        "也愿我们一起，让沉默发声，让平凡不凡。",
        size=11.5, color=INK, after=42, line=1.55,
    )
    p = add_text(doc, "非凡教育  ·  三年育林，百年根基", size=10.5, color=FOREST, bold=True, after=0)
    shade_paragraph(p, PALE_YELLOW, border=YELLOW)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
